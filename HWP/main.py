"""
HWP Studio - Korean Document Editor
A Hangul-compatible document editor for Windows
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser, font
import tkinter.font as tkfont
import json
import os
import sys
import re
from datetime import datetime
from pathlib import Path

# Optional imports
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


# ─────────────────────────────────────────────
# THEME
# ─────────────────────────────────────────────
THEME = {
    "bg":           "#1E1E2E",
    "sidebar":      "#181825",
    "toolbar":      "#242436",
    "canvas_bg":    "#2A2A3E",
    "paper":        "#FAFAF8",
    "accent":       "#89B4FA",
    "accent2":      "#CBA6F7",
    "success":      "#A6E3A1",
    "warning":      "#F9E2AF",
    "danger":       "#F38BA8",
    "text":         "#CDD6F4",
    "text_dim":     "#6C7086",
    "border":       "#313244",
    "hover":        "#313244",
    "selection":    "#45475A",
    "paper_text":   "#1E1E2E",
    "ruler_bg":     "#242436",
    "ruler_fg":     "#6C7086",
    "tab_active":   "#313244",
}

FONTS_KO = ["맑은 고딕", "굴림", "돋움", "나눔고딕", "바탕", "궁서", "Arial Unicode MS"]
FONTS_EN = ["Arial", "Times New Roman", "Calibri", "Georgia", "Courier New"]

FONT_SIZES = [6, 7, 8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 28, 32, 36, 48, 72]

PAGE_SIZES = {
    "A4":     (595, 842),
    "A3":     (842, 1190),
    "Letter": (612, 792),
    "B5":     (499, 708),
}


# ─────────────────────────────────────────────
# MAIN APPLICATION
# ─────────────────────────────────────────────
class HWPStudio:
    def __init__(self, root):
        self.root = root
        self.root.title("HWP Studio – Korean Document Editor")
        self.root.geometry("1280x820")
        self.root.minsize(900, 600)
        self.root.configure(bg=THEME["bg"])

        # State
        self.current_file = None
        self.modified = False
        self.zoom = 100
        self.page_size = "A4"
        self.show_ruler = tk.BooleanVar(value=True)
        self.show_formatting = tk.BooleanVar(value=False)
        self.word_wrap = tk.BooleanVar(value=True)

        # Font state
        self.current_font_family = tk.StringVar(value="맑은 고딕")
        self.current_font_size = tk.StringVar(value="12")
        self.bold_on = tk.BooleanVar(value=False)
        self.italic_on = tk.BooleanVar(value=False)
        self.underline_on = tk.BooleanVar(value=False)
        self.strikethrough_on = tk.BooleanVar(value=False)
        self.align_var = tk.StringVar(value="left")
        self.current_fg = "#1E1E2E"
        self.current_bg_color = "#FFFFFF"

        # Undo/Redo
        self.undo_stack = []
        self.redo_stack = []

        # Tabs (documents)
        self.documents = []   # list of {"title", "content", "file"}
        self.active_doc = 0

        self._build_ui()
        self._bind_shortcuts()
        self._new_document(startup=True)
        self._update_title()

    # ─── BUILD UI ───────────────────────────────
    def _build_ui(self):
        self._build_menu()
        self._build_toolbar()
        self._build_format_bar()
        self._build_main_area()
        self._build_statusbar()

    def _build_menu(self):
        menubar = tk.Menu(self.root, bg=THEME["toolbar"], fg=THEME["text"],
                          activebackground=THEME["accent"], activeforeground="#000",
                          relief="flat", bd=0)
        self.root.config(menu=menubar)

        # File
        file_menu = tk.Menu(menubar, tearoff=0, bg=THEME["toolbar"], fg=THEME["text"],
                            activebackground=THEME["accent"], activeforeground="#000")
        menubar.add_cascade(label="파일(F)", menu=file_menu)
        file_menu.add_command(label="새 문서      Ctrl+N", command=self._new_document)
        file_menu.add_command(label="열기...      Ctrl+O", command=self._open_file)
        file_menu.add_separator()
        file_menu.add_command(label="저장         Ctrl+S", command=self._save_file)
        file_menu.add_command(label="다른 이름으로 저장...", command=self._save_as)
        file_menu.add_separator()
        file_menu.add_command(label="PDF로 내보내기...", command=self._export_pdf)
        file_menu.add_command(label="DOCX로 내보내기...", command=self._export_docx)
        file_menu.add_separator()
        file_menu.add_command(label="종료         Alt+F4", command=self._quit)

        # Edit
        edit_menu = tk.Menu(menubar, tearoff=0, bg=THEME["toolbar"], fg=THEME["text"],
                            activebackground=THEME["accent"], activeforeground="#000")
        menubar.add_cascade(label="편집(E)", menu=edit_menu)
        edit_menu.add_command(label="실행 취소    Ctrl+Z", command=self._undo)
        edit_menu.add_command(label="다시 실행    Ctrl+Y", command=self._redo)
        edit_menu.add_separator()
        edit_menu.add_command(label="잘라내기     Ctrl+X", command=lambda: self.text.event_generate("<<Cut>>"))
        edit_menu.add_command(label="복사         Ctrl+C", command=lambda: self.text.event_generate("<<Copy>>"))
        edit_menu.add_command(label="붙여넣기     Ctrl+V", command=lambda: self.text.event_generate("<<Paste>>"))
        edit_menu.add_command(label="모두 선택    Ctrl+A", command=lambda: self.text.tag_add("sel", "1.0", "end"))
        edit_menu.add_separator()
        edit_menu.add_command(label="찾기/바꾸기  Ctrl+H", command=self._find_replace)

        # Insert
        insert_menu = tk.Menu(menubar, tearoff=0, bg=THEME["toolbar"], fg=THEME["text"],
                              activebackground=THEME["accent"], activeforeground="#000")
        menubar.add_cascade(label="삽입(I)", menu=insert_menu)
        insert_menu.add_command(label="표 삽입...", command=self._insert_table_dialog)
        insert_menu.add_command(label="체크박스 삽입", command=self._insert_checkbox)
        insert_menu.add_command(label="구분선 삽입", command=self._insert_divider)
        insert_menu.add_separator()
        insert_menu.add_command(label="날짜/시간 삽입", command=self._insert_datetime)
        insert_menu.add_command(label="페이지 번호", command=self._insert_page_num)

        # Format
        format_menu = tk.Menu(menubar, tearoff=0, bg=THEME["toolbar"], fg=THEME["text"],
                              activebackground=THEME["accent"], activeforeground="#000")
        menubar.add_cascade(label="서식(O)", menu=format_menu)
        format_menu.add_command(label="글꼴...", command=self._font_dialog)
        format_menu.add_command(label="문단...", command=self._paragraph_dialog)
        format_menu.add_separator()
        format_menu.add_checkbutton(label="눈금자 표시", variable=self.show_ruler,
                                     command=self._toggle_ruler)
        format_menu.add_checkbutton(label="서식 기호 표시", variable=self.show_formatting,
                                     command=self._toggle_formatting_marks)

        # View
        view_menu = tk.Menu(menubar, tearoff=0, bg=THEME["toolbar"], fg=THEME["text"],
                            activebackground=THEME["accent"], activeforeground="#000")
        menubar.add_cascade(label="보기(V)", menu=view_menu)
        for z in [75, 100, 125, 150, 200]:
            view_menu.add_command(label=f"확대/축소 {z}%",
                                  command=lambda z=z: self._set_zoom(z))

        # Help
        help_menu = tk.Menu(menubar, tearoff=0, bg=THEME["toolbar"], fg=THEME["text"],
                            activebackground=THEME["accent"], activeforeground="#000")
        menubar.add_cascade(label="도움말(H)", menu=help_menu)
        help_menu.add_command(label="HWP Studio 정보", command=self._about)

    def _build_toolbar(self):
        self.toolbar = tk.Frame(self.root, bg=THEME["toolbar"], height=44, pady=4)
        self.toolbar.pack(side="top", fill="x")
        self.toolbar.pack_propagate(False)

        btn_cfg = dict(bg=THEME["toolbar"], fg=THEME["text"], relief="flat",
                       bd=0, padx=8, pady=4, cursor="hand2",
                       activebackground=THEME["hover"], activeforeground=THEME["accent"],
                       font=("Segoe UI", 10))

        def tb_btn(parent, text, cmd, tooltip=None):
            b = tk.Button(parent, text=text, command=cmd, **btn_cfg)
            b.pack(side="left", padx=1)
            if tooltip:
                self._add_tooltip(b, tooltip)
            return b

        def sep():
            tk.Frame(self.toolbar, bg=THEME["border"], width=1).pack(side="left", fill="y", padx=4, pady=4)

        tb_btn(self.toolbar, "□ 새문서", self._new_document, "새 문서 (Ctrl+N)")
        tb_btn(self.toolbar, "📂 열기", self._open_file, "파일 열기 (Ctrl+O)")
        tb_btn(self.toolbar, "💾 저장", self._save_file, "저장 (Ctrl+S)")
        sep()
        tb_btn(self.toolbar, "↩ 실행취소", self._undo, "실행 취소 (Ctrl+Z)")
        tb_btn(self.toolbar, "↪ 다시실행", self._redo, "다시 실행 (Ctrl+Y)")
        sep()
        tb_btn(self.toolbar, "✂ 잘라내기", lambda: self.text.event_generate("<<Cut>>"))
        tb_btn(self.toolbar, "📋 복사", lambda: self.text.event_generate("<<Copy>>"))
        tb_btn(self.toolbar, "📌 붙여넣기", lambda: self.text.event_generate("<<Paste>>"))
        sep()
        tb_btn(self.toolbar, "🔍 찾기/바꾸기", self._find_replace)
        sep()
        tb_btn(self.toolbar, "☑ 체크박스", self._insert_checkbox, "체크박스 삽입")
        tb_btn(self.toolbar, "⊞ 표", self._insert_table_dialog, "표 삽입")
        sep()
        tb_btn(self.toolbar, "📄 PDF", self._export_pdf, "PDF로 내보내기")
        tb_btn(self.toolbar, "📝 DOCX", self._export_docx, "DOCX로 내보내기")

        # Zoom on right
        zoom_frame = tk.Frame(self.toolbar, bg=THEME["toolbar"])
        zoom_frame.pack(side="right", padx=8)
        tk.Label(zoom_frame, text="확대:", bg=THEME["toolbar"], fg=THEME["text_dim"],
                 font=("Segoe UI", 9)).pack(side="left")
        self.zoom_var = tk.StringVar(value="100%")
        zoom_cb = ttk.Combobox(zoom_frame, textvariable=self.zoom_var, width=6,
                               values=["75%","100%","125%","150%","200%"])
        zoom_cb.pack(side="left", padx=4)
        zoom_cb.bind("<<ComboboxSelected>>", self._on_zoom_change)

    def _build_format_bar(self):
        self.format_bar = tk.Frame(self.root, bg=THEME["sidebar"], height=36, pady=2)
        self.format_bar.pack(side="top", fill="x")
        self.format_bar.pack_propagate(False)

        # Font family
        all_fonts = sorted(set(FONTS_KO + FONTS_EN + list(tkfont.families())))
        font_cb = ttk.Combobox(self.format_bar, textvariable=self.current_font_family,
                               values=all_fonts, width=20)
        font_cb.pack(side="left", padx=(8,2), pady=3)
        font_cb.bind("<<ComboboxSelected>>", self._apply_font)
        font_cb.bind("<Return>", self._apply_font)

        # Font size
        size_cb = ttk.Combobox(self.format_bar, textvariable=self.current_font_size,
                               values=[str(s) for s in FONT_SIZES], width=5)
        size_cb.pack(side="left", padx=2, pady=3)
        size_cb.bind("<<ComboboxSelected>>", self._apply_font)
        size_cb.bind("<Return>", self._apply_font)

        def sep():
            tk.Frame(self.format_bar, bg=THEME["border"], width=1).pack(
                side="left", fill="y", padx=4, pady=4)

        sep()

        toggle_cfg = dict(relief="flat", bd=0, padx=8, pady=3, cursor="hand2",
                          font=("Segoe UI", 10, "bold"), width=2)

        self.btn_bold = tk.Button(self.format_bar, text="B", command=self._toggle_bold,
                                  bg=THEME["sidebar"], fg=THEME["text"], **toggle_cfg)
        self.btn_bold.pack(side="left", padx=1)

        self.btn_italic = tk.Button(self.format_bar, text="I", command=self._toggle_italic,
                                    bg=THEME["sidebar"], fg=THEME["text"],
                                    font=("Segoe UI", 10, "italic"), **{k:v for k,v in toggle_cfg.items() if k!="font"})
        self.btn_italic.pack(side="left", padx=1)

        self.btn_underline = tk.Button(self.format_bar, text="U", command=self._toggle_underline,
                                       bg=THEME["sidebar"], fg=THEME["text"], **toggle_cfg)
        self.btn_underline.pack(side="left", padx=1)

        self.btn_strike = tk.Button(self.format_bar, text="S̶", command=self._toggle_strike,
                                    bg=THEME["sidebar"], fg=THEME["text"], **toggle_cfg)
        self.btn_strike.pack(side="left", padx=1)

        sep()

        # Alignment
        for sym, val, tip in [("≡L","left","왼쪽 정렬"), ("≡C","center","가운데 정렬"),
                               ("≡R","right","오른쪽 정렬"), ("≡J","justify","양쪽 정렬")]:
            b = tk.Button(self.format_bar, text=sym, width=2,
                          command=lambda v=val: self._set_align(v),
                          bg=THEME["sidebar"], fg=THEME["text"],
                          relief="flat", bd=0, padx=6, pady=3, cursor="hand2",
                          font=("Segoe UI", 10))
            b.pack(side="left", padx=1)
            self._add_tooltip(b, tip)

        sep()

        # Color buttons
        tk.Button(self.format_bar, text="A", fg=self.current_fg, bg=THEME["sidebar"],
                  relief="flat", bd=0, padx=6, pady=3, cursor="hand2",
                  font=("Segoe UI", 10, "bold"), command=self._pick_text_color).pack(side="left", padx=1)

        tk.Button(self.format_bar, text="▌", fg="#FFFF00", bg=THEME["sidebar"],
                  relief="flat", bd=0, padx=6, pady=3, cursor="hand2",
                  font=("Segoe UI", 10), command=self._pick_highlight_color).pack(side="left", padx=1)

        sep()

        # List buttons
        tk.Button(self.format_bar, text="• 목록", bg=THEME["sidebar"], fg=THEME["text"],
                  relief="flat", bd=0, padx=6, pady=3, cursor="hand2",
                  font=("Segoe UI", 9), command=self._insert_bullet).pack(side="left", padx=1)

        tk.Button(self.format_bar, text="1. 번호", bg=THEME["sidebar"], fg=THEME["text"],
                  relief="flat", bd=0, padx=6, pady=3, cursor="hand2",
                  font=("Segoe UI", 9), command=self._insert_numbered).pack(side="left", padx=1)

        sep()

        # Indent
        tk.Button(self.format_bar, text="→ 들여", bg=THEME["sidebar"], fg=THEME["text"],
                  relief="flat", bd=0, padx=6, pady=3, cursor="hand2",
                  font=("Segoe UI", 9), command=self._indent).pack(side="left", padx=1)
        tk.Button(self.format_bar, text="← 내어", bg=THEME["sidebar"], fg=THEME["text"],
                  relief="flat", bd=0, padx=6, pady=3, cursor="hand2",
                  font=("Segoe UI", 9), command=self._outdent).pack(side="left", padx=1)

    def _build_main_area(self):
        # Tab bar
        self.tab_bar = tk.Frame(self.root, bg=THEME["bg"], height=32)
        self.tab_bar.pack(side="top", fill="x")
        self.tab_bar.pack_propagate(False)
        self.tabs_frame = tk.Frame(self.tab_bar, bg=THEME["bg"])
        self.tabs_frame.pack(side="left", fill="both", expand=True)
        tk.Button(self.tab_bar, text="+", bg=THEME["bg"], fg=THEME["accent"],
                  relief="flat", bd=0, padx=10, cursor="hand2",
                  command=self._new_document, font=("Segoe UI", 14)).pack(side="left")

        # Ruler
        self.ruler_frame = tk.Frame(self.root, bg=THEME["ruler_bg"], height=20)
        self.ruler_frame.pack(side="top", fill="x")
        self.ruler_canvas = tk.Canvas(self.ruler_frame, bg=THEME["ruler_bg"],
                                      height=20, highlightthickness=0)
        self.ruler_canvas.pack(fill="x")
        self._draw_ruler()

        # Main paned
        main_pane = tk.PanedWindow(self.root, orient="horizontal",
                                   bg=THEME["bg"], sashwidth=4, sashrelief="flat")
        main_pane.pack(side="top", fill="both", expand=True)

        # Sidebar
        self.sidebar = tk.Frame(main_pane, bg=THEME["sidebar"], width=200)
        main_pane.add(self.sidebar, minsize=0)
        self._build_sidebar()

        # Editor area
        editor_frame = tk.Frame(main_pane, bg=THEME["canvas_bg"])
        main_pane.add(editor_frame, minsize=400)

        # Paper wrapper (scrollable)
        self.canvas_scroll = tk.Canvas(editor_frame, bg=THEME["canvas_bg"],
                                       highlightthickness=0)
        vscroll = tk.Scrollbar(editor_frame, orient="vertical",
                               command=self.canvas_scroll.yview)
        hscroll = tk.Scrollbar(editor_frame, orient="horizontal",
                               command=self.canvas_scroll.xview)
        self.canvas_scroll.configure(yscrollcommand=vscroll.set,
                                     xscrollcommand=hscroll.set)
        vscroll.pack(side="right", fill="y")
        hscroll.pack(side="bottom", fill="x")
        self.canvas_scroll.pack(fill="both", expand=True)

        # Paper frame inside canvas
        self.paper_outer = tk.Frame(self.canvas_scroll, bg=THEME["canvas_bg"])
        self.canvas_window = self.canvas_scroll.create_window(
            (0, 0), window=self.paper_outer, anchor="n")
        self.canvas_scroll.bind("<Configure>", self._on_canvas_configure)
        self.paper_outer.bind("<Configure>", self._on_paper_configure)
        self.canvas_scroll.bind("<MouseWheel>", self._on_mousewheel)

        # Paper shadow + paper
        self.paper_shadow = tk.Frame(self.paper_outer, bg="#111122")
        self.paper_shadow.pack(pady=(30,30), padx=30)

        self.paper = tk.Frame(self.paper_shadow, bg=THEME["paper"],
                              padx=80, pady=60)
        self.paper.pack(padx=3, pady=3)

        # Text widget (the actual editor)
        default_font = tkfont.Font(family="맑은 고딕", size=12)
        self.text = tk.Text(
            self.paper,
            font=default_font,
            bg=THEME["paper"],
            fg=THEME["paper_text"],
            insertbackground=THEME["accent"],
            selectbackground=THEME["accent"],
            selectforeground="#000",
            relief="flat",
            bd=0,
            wrap="word",
            undo=True,
            maxundo=50,
            spacing1=2,
            spacing2=2,
            spacing3=2,
            width=65,
            height=50,
            highlightthickness=0,
            padx=0,
            pady=0,
        )
        self.text.pack(fill="both", expand=True)
        self.text.bind("<<Modified>>", self._on_text_modified)
        self.text.bind("<KeyRelease>", self._on_key_release)
        self.text.bind("<Button-1>", self._on_click)
        self.text.bind("<ButtonRelease-1>", self._on_click)

        # Configure tags
        self._configure_tags()

        # Right-click menu
        self._build_context_menu()

    def _build_sidebar(self):
        tk.Label(self.sidebar, text="문서 구조", bg=THEME["sidebar"],
                 fg=THEME["text_dim"], font=("Segoe UI", 9, "bold"),
                 pady=8).pack(fill="x", padx=8)

        # Page info
        self.page_info = tk.Label(self.sidebar, text="페이지 1 / 1",
                                   bg=THEME["sidebar"], fg=THEME["text"],
                                   font=("Segoe UI", 9))
        self.page_info.pack(fill="x", padx=8, pady=2)

        tk.Frame(self.sidebar, bg=THEME["border"], height=1).pack(fill="x", padx=8, pady=4)

        # Page size selector
        tk.Label(self.sidebar, text="용지 크기", bg=THEME["sidebar"],
                 fg=THEME["text_dim"], font=("Segoe UI", 8)).pack(fill="x", padx=8)
        page_var = tk.StringVar(value="A4")
        page_cb = ttk.Combobox(self.sidebar, textvariable=page_var,
                               values=list(PAGE_SIZES.keys()), width=12)
        page_cb.pack(padx=8, pady=2, anchor="w")

        tk.Frame(self.sidebar, bg=THEME["border"], height=1).pack(fill="x", padx=8, pady=4)

        # Quick insert
        tk.Label(self.sidebar, text="빠른 삽입", bg=THEME["sidebar"],
                 fg=THEME["text_dim"], font=("Segoe UI", 8, "bold")).pack(fill="x", padx=8)

        quick_items = [
            ("☑ 체크박스", self._insert_checkbox),
            ("⊞ 표", self._insert_table_dialog),
            ("─ 구분선", self._insert_divider),
            ("📅 날짜", self._insert_datetime),
            ("• 글머리표", self._insert_bullet),
        ]
        for label, cmd in quick_items:
            b = tk.Button(self.sidebar, text=label, command=cmd,
                          bg=THEME["sidebar"], fg=THEME["text"],
                          relief="flat", bd=0, padx=8, pady=4,
                          anchor="w", cursor="hand2", font=("Segoe UI", 9),
                          activebackground=THEME["hover"])
            b.pack(fill="x", padx=4, pady=1)

        tk.Frame(self.sidebar, bg=THEME["border"], height=1).pack(fill="x", padx=8, pady=4)

        # Statistics
        tk.Label(self.sidebar, text="문서 통계", bg=THEME["sidebar"],
                 fg=THEME["text_dim"], font=("Segoe UI", 8, "bold")).pack(fill="x", padx=8)
        self.stat_label = tk.Label(self.sidebar, text="글자: 0\n단어: 0\n줄: 1",
                                   bg=THEME["sidebar"], fg=THEME["text"],
                                   font=("Segoe UI", 8), justify="left")
        self.stat_label.pack(fill="x", padx=8, pady=2)

    def _build_statusbar(self):
        status = tk.Frame(self.root, bg=THEME["toolbar"], height=24)
        status.pack(side="bottom", fill="x")
        status.pack_propagate(False)

        self.status_left = tk.Label(status, text="준비",
                                    bg=THEME["toolbar"], fg=THEME["text_dim"],
                                    font=("Segoe UI", 8))
        self.status_left.pack(side="left", padx=8)

        self.status_pos = tk.Label(status, text="줄 1, 열 1",
                                   bg=THEME["toolbar"], fg=THEME["text_dim"],
                                   font=("Segoe UI", 8))
        self.status_pos.pack(side="right", padx=8)

        self.status_encode = tk.Label(status, text="UTF-8 | 한국어",
                                      bg=THEME["toolbar"], fg=THEME["text_dim"],
                                      font=("Segoe UI", 8))
        self.status_encode.pack(side="right", padx=8)

        self.status_zoom = tk.Label(status, text="100%",
                                    bg=THEME["toolbar"], fg=THEME["text_dim"],
                                    font=("Segoe UI", 8))
        self.status_zoom.pack(side="right", padx=8)

    def _build_context_menu(self):
        self.ctx_menu = tk.Menu(self.root, tearoff=0,
                                bg=THEME["toolbar"], fg=THEME["text"],
                                activebackground=THEME["accent"], activeforeground="#000")
        self.ctx_menu.add_command(label="잘라내기", command=lambda: self.text.event_generate("<<Cut>>"))
        self.ctx_menu.add_command(label="복사", command=lambda: self.text.event_generate("<<Copy>>"))
        self.ctx_menu.add_command(label="붙여넣기", command=lambda: self.text.event_generate("<<Paste>>"))
        self.ctx_menu.add_separator()
        self.ctx_menu.add_command(label="체크박스 삽입", command=self._insert_checkbox)
        self.ctx_menu.add_command(label="표 삽입", command=self._insert_table_dialog)
        self.ctx_menu.add_command(label="구분선 삽입", command=self._insert_divider)
        self.text.bind("<Button-3>", self._show_context_menu)

    def _configure_tags(self):
        self.text.tag_configure("bold", font=("맑은 고딕", 12, "bold"))
        self.text.tag_configure("italic", font=("맑은 고딕", 12, "italic"))
        self.text.tag_configure("underline", underline=True)
        self.text.tag_configure("strikethrough", overstrike=True)
        self.text.tag_configure("align_left", justify="left")
        self.text.tag_configure("align_center", justify="center")
        self.text.tag_configure("align_right", justify="right")
        self.text.tag_configure("checkbox_done", foreground="#A6E3A1")
        self.text.tag_configure("divider", foreground=THEME["accent"])
        self.text.tag_configure("heading1", font=("맑은 고딕", 20, "bold"))
        self.text.tag_configure("heading2", font=("맑은 고딕", 16, "bold"))
        self.text.tag_configure("heading3", font=("맑은 고딕", 13, "bold"))
        self.text.tag_configure("highlight", background="#FFFF99")
        self.text.tag_configure("find_result", background="#FFD700", foreground="#000")

    # ─── TABS ───────────────────────────────────
    def _new_document(self, startup=False):
        doc = {
            "title": f"새 문서 {len(self.documents)+1}",
            "content": "",
            "file": None,
            "modified": False,
        }
        self.documents.append(doc)
        self._rebuild_tabs()
        self._switch_tab(len(self.documents) - 1)
        if not startup:
            self._status("새 문서가 만들어졌습니다.")

    def _rebuild_tabs(self):
        for w in self.tabs_frame.winfo_children():
            w.destroy()
        for i, doc in enumerate(self.documents):
            title = doc["title"] + (" *" if doc.get("modified") else "")
            bg = THEME["tab_active"] if i == self.active_doc else THEME["bg"]
            fg = THEME["accent"] if i == self.active_doc else THEME["text_dim"]
            f = tk.Frame(self.tabs_frame, bg=bg)
            f.pack(side="left")
            tk.Label(f, text=title, bg=bg, fg=fg,
                     font=("Segoe UI", 9), padx=10, pady=6,
                     cursor="hand2").pack(side="left")
            close = tk.Label(f, text="×", bg=bg, fg=THEME["text_dim"],
                             font=("Segoe UI", 9), padx=4, cursor="hand2")
            close.pack(side="left")
            f.bind("<Button-1>", lambda e, idx=i: self._switch_tab(idx))
            for child in f.winfo_children():
                child.bind("<Button-1>", lambda e, idx=i: self._switch_tab(idx))
            close.bind("<Button-1>", lambda e, idx=i: self._close_tab(idx))

    def _switch_tab(self, idx):
        if self.documents:
            # Save current content
            if hasattr(self, 'text'):
                self.documents[self.active_doc]["content"] = self.text.get("1.0", "end-1c")
        self.active_doc = idx
        if hasattr(self, 'text'):
            self.text.delete("1.0", "end")
            self.text.insert("1.0", self.documents[idx]["content"])
            self.text.edit_reset()
        self._rebuild_tabs()
        self._update_title()

    def _close_tab(self, idx):
        if len(self.documents) == 1:
            self._new_document()
            self.documents.pop(0)
            self.active_doc = 0
        else:
            self.documents.pop(idx)
            self.active_doc = max(0, idx - 1)
        self._rebuild_tabs()
        self._switch_tab(self.active_doc)

    # ─── FILE OPS ───────────────────────────────
    def _open_file(self):
        path = filedialog.askopenfilename(
            filetypes=[
                ("지원 파일", "*.hwps *.txt *.md *.docx"),
                ("HWP Studio", "*.hwps"),
                ("텍스트 파일", "*.txt *.md"),
                ("모든 파일", "*.*"),
            ]
        )
        if not path:
            return
        try:
            ext = Path(path).suffix.lower()
            if ext == ".hwps":
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                content = data.get("content", "")
                title = data.get("title", Path(path).stem)
            elif ext in (".docx",) and DOCX_AVAILABLE:
                doc = DocxDocument(path)
                content = "\n".join(p.text for p in doc.paragraphs)
                title = Path(path).stem
            else:
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()
                title = Path(path).stem

            new_doc = {"title": title, "content": content,
                       "file": path, "modified": False}
            self.documents.append(new_doc)
            self._rebuild_tabs()
            self._switch_tab(len(self.documents) - 1)
            self._status(f"파일을 열었습니다: {path}")
        except Exception as e:
            messagebox.showerror("오류", f"파일을 열 수 없습니다:\n{e}")

    def _save_file(self):
        doc = self.documents[self.active_doc]
        if doc["file"] and doc["file"].endswith(".hwps"):
            self._do_save(doc["file"])
        else:
            self._save_as()

    def _save_as(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".hwps",
            filetypes=[("HWP Studio", "*.hwps"), ("텍스트 파일", "*.txt"), ("모든 파일", "*.*")]
        )
        if path:
            self._do_save(path)

    def _do_save(self, path):
        content = self.text.get("1.0", "end-1c")
        self.documents[self.active_doc]["content"] = content
        doc = self.documents[self.active_doc]
        try:
            if path.endswith(".hwps"):
                data = {
                    "title": doc["title"],
                    "content": content,
                    "saved_at": datetime.now().isoformat(),
                    "version": "1.0"
                }
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
            doc["file"] = path
            doc["title"] = Path(path).stem
            doc["modified"] = False
            self._rebuild_tabs()
            self._update_title()
            self._status(f"저장 완료: {path}")
        except Exception as e:
            messagebox.showerror("저장 오류", str(e))

    def _export_pdf(self):
        if not PDF_AVAILABLE:
            messagebox.showwarning("PDF 내보내기",
                "reportlab이 설치되어 있지 않습니다.\n"
                "pip install reportlab 을 실행하세요.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF 파일", "*.pdf")])
        if not path:
            return
        try:
            content = self.text.get("1.0", "end-1c")
            c = rl_canvas.Canvas(path, pagesize=A4)
            w, h = A4
            y = h - 60
            c.setFont("Helvetica", 11)
            for line in content.split("\n"):
                if y < 60:
                    c.showPage()
                    y = h - 60
                    c.setFont("Helvetica", 11)
                c.drawString(60, y, line[:90])
                y -= 18
            c.save()
            self._status(f"PDF 내보내기 완료: {path}")
            messagebox.showinfo("완료", f"PDF가 저장되었습니다:\n{path}")
        except Exception as e:
            messagebox.showerror("PDF 오류", str(e))

    def _export_docx(self):
        if not DOCX_AVAILABLE:
            messagebox.showwarning("DOCX 내보내기",
                "python-docx가 설치되어 있지 않습니다.\n"
                "pip install python-docx 를 실행하세요.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word 문서", "*.docx")])
        if not path:
            return
        try:
            content = self.text.get("1.0", "end-1c")
            doc = DocxDocument()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(path)
            self._status(f"DOCX 내보내기 완료: {path}")
            messagebox.showinfo("완료", f"DOCX가 저장되었습니다:\n{path}")
        except Exception as e:
            messagebox.showerror("DOCX 오류", str(e))

    # ─── TEXT FORMATTING ────────────────────────
    def _apply_font(self, event=None):
        family = self.current_font_family.get()
        try:
            size = int(self.current_font_size.get())
        except ValueError:
            size = 12
        weight = "bold" if self.bold_on.get() else "normal"
        slant = "italic" if self.italic_on.get() else "roman"
        f = tkfont.Font(family=family, size=size, weight=weight, slant=slant)
        tag = f"font_{family}_{size}_{weight}_{slant}"
        self.text.tag_configure(tag, font=f)
        try:
            sel_start = self.text.index("sel.first")
            sel_end = self.text.index("sel.last")
            self.text.tag_add(tag, sel_start, sel_end)
        except tk.TclError:
            self.text.configure(font=f)

    def _toggle_bold(self):
        self.bold_on.set(not self.bold_on.get())
        self.btn_bold.configure(
            bg=THEME["accent"] if self.bold_on.get() else THEME["sidebar"],
            fg="#000" if self.bold_on.get() else THEME["text"])
        self._apply_tag_to_selection("bold")

    def _toggle_italic(self):
        self.italic_on.set(not self.italic_on.get())
        self.btn_italic.configure(
            bg=THEME["accent"] if self.italic_on.get() else THEME["sidebar"],
            fg="#000" if self.italic_on.get() else THEME["text"])
        self._apply_tag_to_selection("italic")

    def _toggle_underline(self):
        self.underline_on.set(not self.underline_on.get())
        self.btn_underline.configure(
            bg=THEME["accent"] if self.underline_on.get() else THEME["sidebar"],
            fg="#000" if self.underline_on.get() else THEME["text"])
        self._apply_tag_to_selection("underline")

    def _toggle_strike(self):
        self.strikethrough_on.set(not self.strikethrough_on.get())
        self.btn_strike.configure(
            bg=THEME["accent"] if self.strikethrough_on.get() else THEME["sidebar"],
            fg="#000" if self.strikethrough_on.get() else THEME["text"])
        self._apply_tag_to_selection("strikethrough")

    def _apply_tag_to_selection(self, tag):
        try:
            s = self.text.index("sel.first")
            e = self.text.index("sel.last")
            if tag in self.text.tag_names(s):
                self.text.tag_remove(tag, s, e)
            else:
                self.text.tag_add(tag, s, e)
        except tk.TclError:
            pass

    def _set_align(self, align):
        self.align_var.set(align)
        tag = f"align_{align}"
        line_start = self.text.index("insert linestart")
        line_end = self.text.index("insert lineend")
        for a in ["align_left", "align_center", "align_right"]:
            self.text.tag_remove(a, line_start, line_end)
        self.text.tag_add(tag, line_start, line_end)

    def _pick_text_color(self):
        color = colorchooser.askcolor(color=self.current_fg, title="글자 색 선택")[1]
        if color:
            self.current_fg = color
            tag = f"color_{color.replace('#','')}"
            self.text.tag_configure(tag, foreground=color)
            try:
                self.text.tag_add(tag, "sel.first", "sel.last")
            except tk.TclError:
                pass

    def _pick_highlight_color(self):
        color = colorchooser.askcolor(color="#FFFF99", title="형광펜 색 선택")[1]
        if color:
            tag = f"highlight_{color.replace('#','')}"
            self.text.tag_configure(tag, background=color)
            try:
                self.text.tag_add(tag, "sel.first", "sel.last")
            except tk.TclError:
                pass

    # ─── INSERT ─────────────────────────────────
    def _insert_checkbox(self):
        self.text.insert("insert", "☐ ")
        self._status("체크박스가 삽입되었습니다. 클릭하여 선택/해제하세요.")
        # Bind click on checkboxes
        self.text.bind("<Button-1>", self._handle_checkbox_click)

    def _handle_checkbox_click(self, event):
        idx = self.text.index(f"@{event.x},{event.y}")
        # Check char before and at click
        try:
            char = self.text.get(idx)
            if char == "☐":
                self.text.delete(idx)
                self.text.insert(idx, "☑")
            elif char == "☑":
                self.text.delete(idx)
                self.text.insert(idx, "☐")
        except:
            pass
        self._on_click(event)

    def _insert_table_dialog(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("표 삽입")
        dlg.geometry("280x180")
        dlg.configure(bg=THEME["bg"])
        dlg.transient(self.root)
        dlg.grab_set()

        tk.Label(dlg, text="표 삽입", bg=THEME["bg"], fg=THEME["text"],
                 font=("Segoe UI", 12, "bold")).pack(pady=(16,8))

        frame = tk.Frame(dlg, bg=THEME["bg"])
        frame.pack(pady=4)

        tk.Label(frame, text="행:", bg=THEME["bg"], fg=THEME["text"],
                 font=("Segoe UI", 10)).grid(row=0, column=0, padx=8, pady=4)
        rows_var = tk.IntVar(value=3)
        tk.Spinbox(frame, from_=1, to=20, textvariable=rows_var, width=6,
                   bg=THEME["toolbar"], fg=THEME["text"]).grid(row=0, column=1)

        tk.Label(frame, text="열:", bg=THEME["bg"], fg=THEME["text"],
                 font=("Segoe UI", 10)).grid(row=1, column=0, padx=8, pady=4)
        cols_var = tk.IntVar(value=3)
        tk.Spinbox(frame, from_=1, to=10, textvariable=cols_var, width=6,
                   bg=THEME["toolbar"], fg=THEME["text"]).grid(row=1, column=1)

        def insert():
            r, c = rows_var.get(), cols_var.get()
            dlg.destroy()
            self._insert_table(r, c)

        tk.Button(dlg, text="삽입", command=insert,
                  bg=THEME["accent"], fg="#000",
                  relief="flat", padx=16, pady=6,
                  font=("Segoe UI", 10)).pack(pady=8)

    def _insert_table(self, rows, cols):
        self.text.insert("insert", "\n")
        col_w = 10
        border = "+" + ("-" * col_w + "+") * cols
        self.text.insert("insert", border + "\n")
        for r in range(rows):
            row_str = "|" + (" " * col_w + "|") * cols
            self.text.insert("insert", row_str + "\n")
            self.text.insert("insert", border + "\n")
        self.text.insert("insert", "\n")
        self._status(f"{rows}×{cols} 표가 삽입되었습니다.")

    def _insert_divider(self):
        self.text.insert("insert", "\n" + "─" * 60 + "\n")

    def _insert_datetime(self):
        now = datetime.now().strftime("%Y년 %m월 %d일 %H:%M")
        self.text.insert("insert", now)

    def _insert_page_num(self):
        self.text.insert("insert", "[페이지]")

    def _insert_bullet(self):
        pos = self.text.index("insert linestart")
        self.text.insert(pos, "• ")

    def _insert_numbered(self):
        # Count existing numbered lines
        content = self.text.get("1.0", "insert")
        num = content.count("\n") + 1
        pos = self.text.index("insert linestart")
        self.text.insert(pos, f"{num}. ")

    def _indent(self):
        pos = self.text.index("insert linestart")
        self.text.insert(pos, "    ")

    def _outdent(self):
        pos = self.text.index("insert linestart")
        line = self.text.get(pos, pos + "+4c")
        if line.startswith("    "):
            self.text.delete(pos, pos + "+4c")
        elif line.startswith("\t"):
            self.text.delete(pos, pos + "+1c")

    # ─── FIND & REPLACE ─────────────────────────
    def _find_replace(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("찾기/바꾸기")
        dlg.geometry("380x180")
        dlg.configure(bg=THEME["bg"])
        dlg.transient(self.root)

        tk.Label(dlg, text="찾기:", bg=THEME["bg"], fg=THEME["text"],
                 font=("Segoe UI", 10)).grid(row=0, column=0, padx=12, pady=8, sticky="w")
        find_var = tk.StringVar()
        find_entry = tk.Entry(dlg, textvariable=find_var, width=24,
                              bg=THEME["toolbar"], fg=THEME["text"],
                              insertbackground=THEME["accent"], relief="flat")
        find_entry.grid(row=0, column=1, columnspan=2, padx=8, pady=8)
        find_entry.focus()

        tk.Label(dlg, text="바꾸기:", bg=THEME["bg"], fg=THEME["text"],
                 font=("Segoe UI", 10)).grid(row=1, column=0, padx=12, sticky="w")
        repl_var = tk.StringVar()
        tk.Entry(dlg, textvariable=repl_var, width=24,
                 bg=THEME["toolbar"], fg=THEME["text"],
                 insertbackground=THEME["accent"], relief="flat").grid(
                     row=1, column=1, columnspan=2, padx=8)

        result_label = tk.Label(dlg, text="", bg=THEME["bg"], fg=THEME["accent"],
                                font=("Segoe UI", 9))
        result_label.grid(row=2, column=0, columnspan=3, pady=4)

        def do_find():
            self.text.tag_remove("find_result", "1.0", "end")
            term = find_var.get()
            if not term:
                return
            count = 0
            start = "1.0"
            while True:
                pos = self.text.search(term, start, stopindex="end")
                if not pos:
                    break
                end = f"{pos}+{len(term)}c"
                self.text.tag_add("find_result", pos, end)
                start = end
                count += 1
            result_label.config(text=f"{count}개 찾음")

        def do_replace_all():
            term = find_var.get()
            repl = repl_var.get()
            if not term:
                return
            content = self.text.get("1.0", "end-1c")
            count = content.count(term)
            new_content = content.replace(term, repl)
            self.text.delete("1.0", "end")
            self.text.insert("1.0", new_content)
            result_label.config(text=f"{count}개 바꿈")

        btn_cfg = dict(bg=THEME["accent"], fg="#000", relief="flat",
                       padx=10, pady=4, font=("Segoe UI", 9), cursor="hand2")
        tk.Button(dlg, text="찾기", command=do_find, **btn_cfg).grid(
            row=3, column=0, padx=8, pady=8)
        tk.Button(dlg, text="모두 바꾸기", command=do_replace_all, **btn_cfg).grid(
            row=3, column=1, padx=4)
        tk.Button(dlg, text="닫기", command=dlg.destroy,
                  bg=THEME["toolbar"], fg=THEME["text"],
                  relief="flat", padx=10, pady=4,
                  font=("Segoe UI", 9)).grid(row=3, column=2, padx=8)

    # ─── FONT DIALOG ────────────────────────────
    def _font_dialog(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("글꼴")
        dlg.geometry("420x320")
        dlg.configure(bg=THEME["bg"])
        dlg.transient(self.root)
        dlg.grab_set()

        tk.Label(dlg, text="글꼴 설정", bg=THEME["bg"], fg=THEME["text"],
                 font=("Segoe UI", 13, "bold")).pack(pady=(16,8))

        f = tk.Frame(dlg, bg=THEME["bg"])
        f.pack(fill="x", padx=20)

        tk.Label(f, text="글꼴:", bg=THEME["bg"], fg=THEME["text"]).grid(
            row=0, column=0, sticky="w", pady=4)
        families = sorted(tkfont.families())
        fam_var = tk.StringVar(value=self.current_font_family.get())
        fam_cb = ttk.Combobox(f, textvariable=fam_var, values=families, width=22)
        fam_cb.grid(row=0, column=1, padx=8)

        tk.Label(f, text="크기:", bg=THEME["bg"], fg=THEME["text"]).grid(
            row=1, column=0, sticky="w", pady=4)
        size_var = tk.StringVar(value=self.current_font_size.get())
        ttk.Combobox(f, textvariable=size_var,
                     values=[str(s) for s in FONT_SIZES], width=8).grid(row=1, column=1, padx=8, sticky="w")

        # Style checkboxes
        bold_v = tk.BooleanVar(value=self.bold_on.get())
        italic_v = tk.BooleanVar(value=self.italic_on.get())
        under_v = tk.BooleanVar(value=self.underline_on.get())
        strike_v = tk.BooleanVar(value=self.strikethrough_on.get())

        style_frame = tk.Frame(dlg, bg=THEME["bg"])
        style_frame.pack(fill="x", padx=20, pady=8)
        for text, var in [("굵게", bold_v), ("기울임", italic_v),
                          ("밑줄", under_v), ("취소선", strike_v)]:
            tk.Checkbutton(style_frame, text=text, variable=var,
                           bg=THEME["bg"], fg=THEME["text"],
                           selectcolor=THEME["toolbar"],
                           activebackground=THEME["bg"]).pack(side="left", padx=8)

        # Preview
        preview_frame = tk.Frame(dlg, bg=THEME["paper"], height=50)
        preview_frame.pack(fill="x", padx=20, pady=8)
        preview_label = tk.Label(preview_frame, text="가나다라 ABC 123",
                                 bg=THEME["paper"], fg=THEME["paper_text"])
        preview_label.pack(expand=True)

        def update_preview(*args):
            try:
                sz = int(size_var.get())
            except ValueError:
                sz = 12
            weight = "bold" if bold_v.get() else "normal"
            slant = "italic" if italic_v.get() else "roman"
            preview_label.configure(
                font=tkfont.Font(family=fam_var.get(), size=sz,
                                 weight=weight, slant=slant))

        fam_var.trace("w", update_preview)
        size_var.trace("w", update_preview)
        bold_v.trace("w", update_preview)
        italic_v.trace("w", update_preview)

        def apply():
            self.current_font_family.set(fam_var.get())
            self.current_font_size.set(size_var.get())
            self.bold_on.set(bold_v.get())
            self.italic_on.set(italic_v.get())
            self.underline_on.set(under_v.get())
            self.strikethrough_on.set(strike_v.get())
            self._apply_font()
            dlg.destroy()

        tk.Button(dlg, text="적용", command=apply,
                  bg=THEME["accent"], fg="#000", relief="flat",
                  padx=20, pady=6, font=("Segoe UI", 10)).pack(pady=8)

    def _paragraph_dialog(self):
        messagebox.showinfo("문단", "문단 설정:\n들여쓰기 및 줄 간격은 서식 도구모음을 사용하세요.")

    # ─── ZOOM & RULER ───────────────────────────
    def _set_zoom(self, z):
        self.zoom = z
        self.zoom_var.set(f"{z}%")
        self.status_zoom.configure(text=f"{z}%")
        try:
            size = max(6, int(12 * z / 100))
            self.text.configure(font=tkfont.Font(
                family=self.current_font_family.get(), size=size))
        except:
            pass

    def _on_zoom_change(self, event=None):
        val = self.zoom_var.get().replace("%", "")
        try:
            self._set_zoom(int(val))
        except ValueError:
            pass

    def _toggle_ruler(self):
        if self.show_ruler.get():
            self.ruler_frame.pack(side="top", fill="x", before=self.ruler_frame.master.winfo_children()[0])
        else:
            self.ruler_frame.pack_forget()

    def _toggle_formatting_marks(self):
        pass  # Future: show ¶ markers

    def _draw_ruler(self):
        self.ruler_canvas.update_idletasks()
        w = self.ruler_canvas.winfo_width() or 800
        self.ruler_canvas.delete("all")
        self.ruler_canvas.configure(bg=THEME["ruler_bg"])
        for i in range(0, w, 10):
            h = 10 if i % 50 == 0 else (6 if i % 25 == 0 else 3)
            self.ruler_canvas.create_line(i, 20-h, i, 20,
                                          fill=THEME["ruler_fg"])
            if i % 50 == 0 and i > 0:
                self.ruler_canvas.create_text(i, 5, text=str(i//10),
                                              fill=THEME["ruler_fg"],
                                              font=("Segoe UI", 6))

    # ─── EVENTS ─────────────────────────────────
    def _on_canvas_configure(self, event):
        self.canvas_scroll.itemconfig(self.canvas_window, width=max(event.width, 700))

    def _on_paper_configure(self, event):
        self.canvas_scroll.configure(scrollregion=self.canvas_scroll.bbox("all"))

    def _on_mousewheel(self, event):
        self.canvas_scroll.yview_scroll(int(-1*(event.delta/120)), "units")

    def _on_text_modified(self, event=None):
        if self.text.edit_modified():
            self.documents[self.active_doc]["modified"] = True
            self._rebuild_tabs()
            self._update_stats()
            self.text.edit_modified(False)

    def _on_key_release(self, event=None):
        self._update_cursor_pos()
        self._update_stats()

    def _on_click(self, event=None):
        self._update_cursor_pos()

    def _update_cursor_pos(self):
        try:
            pos = self.text.index("insert")
            line, col = pos.split(".")
            self.status_pos.configure(text=f"줄 {line}, 열 {int(col)+1}")
        except:
            pass

    def _update_stats(self):
        content = self.text.get("1.0", "end-1c")
        chars = len(content)
        words = len(content.split())
        lines = content.count("\n") + 1
        self.stat_label.configure(text=f"글자: {chars}\n단어: {words}\n줄: {lines}")

    def _show_context_menu(self, event):
        self.ctx_menu.tk_popup(event.x_root, event.y_root)

    # ─── UNDO/REDO ──────────────────────────────
    def _undo(self):
        try:
            self.text.edit_undo()
        except tk.TclError:
            self._status("더 이상 실행 취소할 수 없습니다.")

    def _redo(self):
        try:
            self.text.edit_redo()
        except tk.TclError:
            self._status("더 이상 다시 실행할 수 없습니다.")

    # ─── KEYBOARD SHORTCUTS ─────────────────────
    def _bind_shortcuts(self):
        self.root.bind("<Control-n>", lambda e: self._new_document())
        self.root.bind("<Control-o>", lambda e: self._open_file())
        self.root.bind("<Control-s>", lambda e: self._save_file())
        self.root.bind("<Control-z>", lambda e: self._undo())
        self.root.bind("<Control-y>", lambda e: self._redo())
        self.root.bind("<Control-h>", lambda e: self._find_replace())
        self.root.bind("<Control-b>", lambda e: self._toggle_bold())
        self.root.bind("<Control-i>", lambda e: self._toggle_italic())
        self.root.bind("<Control-u>", lambda e: self._toggle_underline())
        self.root.bind("<Control-equal>", lambda e: self._set_zoom(min(200, self.zoom+25)))
        self.root.bind("<Control-minus>", lambda e: self._set_zoom(max(50, self.zoom-25)))
        self.root.bind("<Alt-F4>", lambda e: self._quit())
        self.root.bind("<F1>", lambda e: self._about())

    # ─── HELPERS ────────────────────────────────
    def _add_tooltip(self, widget, text):
        def show(e):
            tip = tk.Toplevel(widget)
            tip.wm_overrideredirect(True)
            tip.wm_geometry(f"+{e.x_root+10}+{e.y_root+20}")
            tk.Label(tip, text=text, bg="#FFFFE0", fg="#000",
                     relief="solid", bd=1, padx=4, pady=2,
                     font=("Segoe UI", 8)).pack()
            widget._tip = tip
        def hide(e):
            if hasattr(widget, "_tip"):
                widget._tip.destroy()
                del widget._tip
        widget.bind("<Enter>", show)
        widget.bind("<Leave>", hide)

    def _update_title(self):
        doc = self.documents[self.active_doc] if self.documents else None
        title = doc["title"] if doc else "HWP Studio"
        self.root.title(f"HWP Studio – {title}")

    def _status(self, msg):
        self.status_left.configure(text=msg)
        self.root.after(4000, lambda: self.status_left.configure(text="준비"))

    def _about(self):
        messagebox.showinfo("HWP Studio 정보",
            "HWP Studio v1.0\n\n"
            "한국어 문서 편집기\n"
            "Python + Tkinter 기반\n\n"
            "지원 형식:\n"
            "  • .hwps (기본 형식)\n"
            "  • .txt / .md 열기\n"
            "  • .docx 내보내기\n"
            "  • .pdf 내보내기\n\n"
            "단축키:\n"
            "  Ctrl+N: 새 문서\n"
            "  Ctrl+O: 열기\n"
            "  Ctrl+S: 저장\n"
            "  Ctrl+B/I/U: 굵게/기울임/밑줄\n"
            "  Ctrl+H: 찾기/바꾸기\n"
            "  Ctrl+Z/Y: 실행취소/다시실행")

    def _quit(self):
        # Check unsaved
        unsaved = [d["title"] for d in self.documents if d.get("modified")]
        if unsaved:
            ans = messagebox.askyesnocancel(
                "종료",
                f"저장하지 않은 문서가 있습니다:\n{', '.join(unsaved)}\n\n저장하고 종료하시겠습니까?")
            if ans is None:
                return
            if ans:
                self._save_file()
        self.root.destroy()


# ─────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────
def main():
    root = tk.Tk()

    # DPI awareness (Windows)
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except:
        pass

    # Style
    style = ttk.Style()
    try:
        style.theme_use("clam")
    except:
        pass
    style.configure("TCombobox",
                    fieldbackground=THEME["toolbar"],
                    background=THEME["toolbar"],
                    foreground=THEME["text"],
                    selectbackground=THEME["accent"],
                    selectforeground="#000")

    app = HWPStudio(root)
    root.protocol("WM_DELETE_WINDOW", app._quit)
    root.mainloop()


if __name__ == "__main__":
    main()
